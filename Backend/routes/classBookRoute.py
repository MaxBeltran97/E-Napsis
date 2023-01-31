import flask
import os
import pythoncom
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx2pdf import convert
from docx import Document
from docx.shared import Cm, Inches
from flask import request, Flask, send_file
from models.calendarCourse import *
from models.course import *
from models.courseTeller import *
from models.teller import *
from models.participant import *
from models.company import *
from models.logo import *
from models.calendarCourseAttendance import *
from models.courseActvityContentHours import *
from models.calendarCourseEvaluation import *





app = Flask(__name__)



UPLOAD_FOLDER_LOGOS = 'assets/subCompanyLogos'
UPLOAD_FOLDER_PATH = 'assets/classBook'
app.config['UPLOAD_FOLDER_LOGOS'] = UPLOAD_FOLDER_LOGOS
app.config['UPLOAD_FOLDER_PATH'] = UPLOAD_FOLDER_PATH

classBook = flask.Blueprint('classBook', __name__)


@classBook.route('/api/classBook/<_id>', methods=['GET'])
def create_classbook(_id):
    try:
 
        calendar = CalendarCourse.query.get(_id)
        course = Course.query.filter_by(_id = calendar.course_id)
        teller = CourseTeller.query.filter_by(course_id = calendar.course_id)
        logo = Logo.query.get(calendar.logo_id)
        cont = 0
    
        for i in course:
            activityName = i.activityName
            sence = i.sence
        
        tellers = ''
        for i in teller:
            relator = Teller.query.get(i.teller_id)
            tellers += f'{relator.fullName} {relator.lastName} {relator.motherLastName} / '
            
        fechai = calendar.startDate.strftime("%d-%m-%Y")
        fechaf = calendar.endDate.strftime("%d-%m-%Y")
        

        template2 = [f'\nCODIGO INTERNO CURSO\t\t\t\t: {calendar.internalCode}',f'NOMBRE OTEC\t\t\t\t\t: {logo.title}',f'NOMBRE ACTIVIDAD\t\t\t\t\t: {activityName}',f'CODIGO AUTORIZADO POR SENCE\t\t\t: {sence}','N REGISTRO SENCE\t\t\t\t\t: ',f'FECHA DE INICIO\t\t\t\t\t: {fechai}', f'FECHA DE TÉRMINO\t\t\t\t\t: {fechaf}',
        f"LUGAR DE EJECUCIÓN\t\t\t\t\t: {calendar.ejecutionPlace}", f"DURACIÓN\t\t\t\t\t\t: {calendar.courseTotalHours}", f"NOMBRE(S) RELATOR(ES)\t\t\t\t: {tellers}", 'OBSERVACIONES\t\t\t\t\t: '
        ]

        # Crear un objeto Document y agregar los parrafos modificados
        doc = Document()
       
       

        section = doc.sections[-1]
        new_width, new_height = section.page_height, section.page_width
        new_width, new_height = Cm(21.59), Cm(35.56)
        section.page_width = new_height
        section.page_height = new_width
        section.orientation = WD_ORIENT.LANDSCAPE
        
         # Se agregar el título
        heading = doc.add_heading('LIBRO DE CONTROL DE CLASES', 1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

         #Añadir imagen
        path = os.path.join(
            app.config["UPLOAD_FOLDER_LOGOS"], logo.logo_img
        )


        doc.add_picture(path, height=Cm(2))


        for item in template2:
            doc.add_paragraph(item)

        
        doc.add_page_break()


        heading = doc.add_heading('ANTECEDENTES PARTICIPANTES', 1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        

        table = doc.add_table(rows=1, cols=7)
        
        # Agregar los encabezados de las columnas
        header_cells = table.rows[0].cells
        header_cells[0].text = "N°"
        header_cells[1].text = "APELLIDOS,NOMBRE"
        header_cells[2].text = "RUT"
        header_cells[3].text = "EMAIL"
        header_cells[4].text = "EMPRESA"
        header_cells[5].text = "CARGO DESEMPEÑADO"
        header_cells[6].text = "FIRMA"


        participants = Participant.query.filter_by(calendarCourse_id = calendar._id)
        
        
        
        for i in participants:
            if cont % 14 == 0 and cont != 0:
                
                row = table.add_row()
                row.cells[0].text = "N°"
                row.cells[1].text = "APELLIDOS,NOMBRE"
                row.cells[2].text = "RUT"
                row.cells[3].text = "EMAIL"
                row.cells[4].text = "EMPRESA"
                row.cells[5].text = "CARGO DESEMPEÑADO"
                row.cells[6].text = "FIRMA"

            name = i.fullName + ' ' + i.lastName + ' ' + i.motherLastName
            cont += 1
            row = table.add_row()
            row.cells[0].text = str(cont)
            row.cells[1].text = name
            row.cells[2].text = i.rut
            row.cells[3].text = i.email
            company = None
            
            if i.company_id == None:
                row.cells[4].text = 'Particular'
            else:
                company = Company.query.get(i.company_id)
                row.cells[4].text = company.fantasyName

            row.cells[5].text = i.position

        table.style = 'Table Grid'
        
        
        for row in table.rows:
            row.height = Cm(1)
            
        
        for cell in table.columns[0].cells:
            cell.width = Inches(0.5)
            
        doc.add_page_break()

        
       
        #table = doc.add_table(rows=0, cols=10)

        data = ['N°', 'APELLIDOS,NOMBRE']

        attendances = CalendarCourseAttendance.query.filter_by(calendarCourse_id = calendar._id)

        contI = 0
        contT = 3
        
        
        for i in attendances:
            
            
            if contI <= contT:
                
                data.append(i.date.strftime("%d-%m-%Y"))
                data.append('Firma')
                
            if contT == contI or contI == attendances.count():
                contP = 0
                heading = doc.add_heading('CONTROL DE ASISTENCIA DE PARTICIPANTES', 1)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                table = doc.add_table(rows=0, cols=10)
                
                

                for j in participants:
                    if contP % 14 == 0:
                        row = table.add_row()
                        for k in range(len(data)):
                            row.cells[k].text = data[k]
                            
                    row = table.add_row()
                        
                    name = j.fullName + ' ' + j.lastName + ' ' + j.motherLastName
                    contP += 1
                    row.cells[0].text =str(contP)
                    row.cells[1].text = name

                
                contT += 4
                
                data = ['N°', 'APELLIDOS,NOMBRE']
                doc.add_page_break()
                
            contI += 1
            table.style = 'Table Grid'
            for row in table.rows:
                row.height = Cm(1)
            
        
            for cell in table.columns[0].cells:
                cell.width = Inches(0.5)

        heading = doc.add_heading('CONTENIDOS', 1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        table = doc.add_table(rows=1, cols=6)

        header_cells = table.rows[0].cells
        header_cells[0].text = "FECHA"
        header_cells[1].text = "TEMAS"
        header_cells[2].text = "ACTIVIDADES"
        header_cells[3].text = "HORA INICIO"
        header_cells[4].text = "HORA TÉRMINO"
        header_cells[5].text = "FIRMA INSTRUCTOR"
      
        contenido = CourseActivityContentHours.query.filter_by(course_id=calendar.course_id)

        for i in contenido:
            row = table.add_row()
            row.cells[1].text = i.activity
            row.cells[2].text = i.content 
        
        table.style = 'Table Grid'
        flag = True
        for row in table.rows:
            if flag:
                flag = False
                continue
            row.height = Cm(3)

        for cell in table.columns[1].cells:
            cell.width = Inches(3)

        for cell in table.columns[2].cells:
            cell.width = Inches(3)
    
        for cell in table.columns[3].cells:
            cell.width = Inches(0.5)

        for cell in table.columns[4].cells:
            cell.width = Inches(0.5)


        doc.add_page_break()


        data = ['N°', 'APELLIDOS,NOMBRE']

        evaluations = CalendarCourseEvaluation.query.filter_by(calendarCourse_id = calendar._id)
        
        
        heading = doc.add_heading('CONTROL DE EVALUACIONES', 1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table = doc.add_table(rows=0, cols=evaluations.count()+4)

        for i in evaluations:
            date = i.evaluationDate.strftime("%d-%m-%Y")
            data.append(f"{date}\n {i.percentage}%")

        data.append('NOTA FINAL')
        data.append('FIRMA')
        
        cont = 0

        for i in participants:
            if cont % 14 == 0:
                row = table.add_row()
                for k in range(len(data)):
                    row.cells[k].text = data[k]
            row = table.add_row()
            cont += 1
            name = j.fullName + ' ' + j.lastName + ' ' + j.motherLastName
            row.cells[0].text =str(cont)
            row.cells[1].text = name
        
        table.style = 'Table Grid'
        
        for row in table.rows:
            row.height = Cm(1)
            
    
        for cell in table.columns[0].cells:
            cell.width = Inches(0.5)

        

        
        pathWord = os.path.join(
            app.config['UPLOAD_FOLDER_PATH'], 'documento.docx'
        )
        pathPDF = os.path.join(
            app.config['UPLOAD_FOLDER_PATH'], 'document.pdf'
        )
        # Guardar el objeto Document como un archivo .docx
        doc.save(pathWord)
        pythoncom.CoInitialize()
        convert(pathWord, pathPDF)
        
        return send_file(pathPDF, as_attachment=True)
        
            
    except Exception as e:
        print(e)
        return {
            "ok": False
        }, 500
   







